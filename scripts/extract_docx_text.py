#!/usr/bin/env python3
"""Extract text from Word (.docx) files.

Uses python-docx — pure Python, no system deps. Only handles the modern
.docx (OOXML) format; legacy .doc (binary) is not supported.

Usage:
    python3 scripts/extract_docx_text.py <path_to_docx>

Output:
    JSON with status, extracted text (paragraphs + table cells), and
    paragraph/table counts.
"""

from __future__ import annotations

import json
import logging
import sys
from pathlib import Path

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger("extract_docx_text")

MAX_CHARS = 500_000  # safety limit, mirrors extract_pdf_text.py's per-page cap


def _table_to_text(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def extract_docx_text(path: str) -> dict:
    p = Path(path)
    if not p.is_file():
        return {"status": "error", "message": f"File not found: {path}"}
    if p.stat().st_size == 0:
        return {"status": "error", "message": "File is empty"}
    if p.stat().st_size > 100 * 1024 * 1024:
        return {"status": "error", "message": "File too large (>100 MB)"}

    try:
        import docx
    except ImportError:
        return {
            "status": "error",
            "message": "python-docx not available. Install: pip install python-docx",
        }

    try:
        document = docx.Document(path)
    except Exception as e:
        msg = str(e).lower()
        if "not a zip file" in msg or "package not found" in msg:
            return {
                "status": "error",
                "message": (
                    "File is not a valid .docx (OOXML) file — legacy binary "
                    ".doc format is not supported."
                ),
                "reason": "unsupported_legacy_doc",
            }
        return {"status": "error", "message": f"Failed to open document: {e}"}

    try:
        paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
        table_texts = [_table_to_text(t) for t in document.tables]
        parts = paragraphs + table_texts
        text = "\n\n".join(parts)
        if len(text) > MAX_CHARS:
            text = text[:MAX_CHARS]
            truncated = True
        else:
            truncated = False

        return {
            "status": "success",
            "method": "python-docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(document.tables),
            "total_chars": len(text),
            "truncated": truncated,
            "text": text,
        }
    except Exception as e:
        return {"status": "error", "message": f"Text extraction failed: {e}"}


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(
            json.dumps(
                {"status": "error", "message": "Usage: extract_docx_text.py <path>"}
            )
        )
        sys.exit(1)

    result = extract_docx_text(sys.argv[1])
    print(json.dumps(result, indent=2, ensure_ascii=False))
